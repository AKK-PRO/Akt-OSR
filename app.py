from flask import Flask, render_template, request, redirect, session, url_for, send_file, send_from_directory
from docx import Document
from docx.shared import Mm
from io import BytesIO
from werkzeug.security import check_password_hash, generate_password_hash
from flask_session import Session
from PIL import Image
import base64
import psycopg2
import os
import datetime

app = Flask(__name__)
app.secret_key = "supersecretkey"
app.config["SESSION_TYPE"] = "filesystem"
Session(app)

# PostgreSQL config
DB_CONFIG = {
    "host": "dpg-d1bto73e5dus73evbl80-a",
    "dbname": "aktdb",
    "user": "aktuser",
    "password": "Z7MAXy7n6KUtc3EcbwsY8JevdBINASlc",
    "port": 5432
}

def get_connection():
    return psycopg2.connect(
        host=DB_CONFIG["host"],
        dbname=DB_CONFIG["dbname"],
        user=DB_CONFIG["user"],
        password=DB_CONFIG["password"],
        port=DB_CONFIG["port"]
    )

# Пользователи
users = {
    "admin": generate_password_hash("1"),
    "worker": generate_password_hash("2")
}

def save_signature_base64(data_url, path):
    header, encoded = data_url.split(",", 1)
    binary_data = base64.b64decode(encoded)
    image = Image.open(BytesIO(binary_data))
    image.save(path, "PNG")

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/login/<role>", methods=["GET", "POST"])
def login(role):
    role_name = "Администратор" if role == "admin" else "Сотрудник"
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]
        if username in users and check_password_hash(users[username], password):
            session["user"] = username
            session["role"] = role
            return redirect(url_for("form"))
        else:
            return render_template("login.html", error="Неверный логин или пароль", role_name=role_name)
    return render_template("login.html", role_name=role_name)

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("index"))

def init_db():
    conn = get_connection()
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS akt_history (
            id SERIAL PRIMARY KEY,
            created_by TEXT,
            akt_number TEXT,
            akt_date TEXT,
            object_description TEXT,
            contractor_name TEXT,
            created_at TEXT
        );
    """)
    conn.commit()
    cur.close()
    conn.close()

def save_to_db(data):
    conn = get_connection()
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO akt_history (
            created_by, akt_number, akt_date, object_description, contractor_name, created_at
        ) VALUES (%s, %s, %s, %s, %s, %s)
        RETURNING id;
    """, (
        data["created_by"],
        data["akt_number"],
        data["akt_date"],
        data["object_description"],
        data["contractor_name"],
        datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    ))
    new_id = cur.fetchone()[0]
    conn.commit()
    cur.close()
    conn.close()
    return new_id

def get_history():
    conn = get_connection()
    cur = conn.cursor()
    cur.execute("SELECT * FROM akt_history ORDER BY id DESC")
    result = cur.fetchall()
    cur.close()
    conn.close()
    return result

def replace_text_preserve_style(paragraph, mapping):
    for run in paragraph.runs:
        for key, value in mapping.items():
            if key in run.text:
                run.text = run.text.replace(key, value)

@app.route("/form", methods=["GET", "POST"])
def form():
    if "user" not in session:
        return redirect(url_for("index"))

    if request.method == "POST":
        fields = {f"{{{field}}}": request.form.get(field, "") for field in [
            "akt_number", "akt_date", "object_description", "contractor_name",
            "contractor_rep", "tech_rep", "author_rep", "additional_rep",
            "work_description", "project_docs", "materials", "proof",
            "deviations", "start_date", "end_date", "next_work"
        ]}

        akt_number = request.form.get("akt_number", "akt")
        tech_data = request.form.get("signature_tech_data")
        author_data = request.form.get("signature_author_data")

        tech_path = f"signatures/tech_{akt_number}.png"
        author_path = f"signatures/author_{akt_number}.png"
        os.makedirs("signatures", exist_ok=True)

        if tech_data:
            save_signature_base64(tech_data, tech_path)
        if author_data:
            save_signature_base64(author_data, author_path)

        doc = Document("template.docx")

        # Замена текстовых плейсхолдеров
        for paragraph in doc.paragraphs:
            replace_text_preserve_style(paragraph, fields)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_preserve_style(paragraph, fields)

        # Вставка подписей в обычных абзацах
        for p in doc.paragraphs:
            if "{signature_tech}" in p.text:
                p.clear()
                run = p.add_run()
                run.add_picture(tech_path, width=Mm(40))
            if "{signature_author}" in p.text:
                p.clear()
                run = p.add_run()
                run.add_picture(author_path, width=Mm(40))

        # Вставка подписей в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "{signature_tech}" in p.text:
                            p.clear()
                            run = p.add_run()
                            run.add_picture(tech_path, width=Mm(40))
                        if "{signature_author}" in p.text:
                            p.clear()
                            run = p.add_run()
                            run.add_picture(author_path, width=Mm(40))

        # Сохраняем акт и возвращаем пользователю
        doc_id = save_to_db({
            "created_by": session["user"],
            "akt_number": request.form.get("akt_number", ""),
            "akt_date": request.form.get("akt_date", ""),
            "object_description": request.form.get("object_description", ""),
            "contractor_name": request.form.get("contractor_name", "")
        })

        filename = f"akt_{doc_id}.docx"
        save_path = os.path.join("saved_docs", filename)
        os.makedirs("saved_docs", exist_ok=True)
        doc.save(save_path)

        return send_file(save_path, as_attachment=True)

    return render_template("form.html", username=session['user'], role=session['role'])

@app.route("/history")
def history():
    if "user" not in session:
        return redirect(url_for("login"))
    rows = get_history()
    return render_template("history.html", rows=rows)

@app.route("/download/<filename>")
def download(filename):
    return send_from_directory("saved_docs", filename, as_attachment=True)

if __name__ == "__main__":
    init_db()
    app.run(host="0.0.0.0", port=10000)
